import os
from docx import Document as DocxDocument
from langchain_community.document_loaders import TextLoader
from langchain.schema import Document
from .ocr import load_docs_with_ocr  # handles both PDFs & images

def load_all_documents(dir_path: str) -> list[Document]:
    docs: list[Document] = []

    for fname in os.listdir(dir_path):
        path = os.path.join(dir_path, fname)
        if not os.path.isfile(path):
            continue

        ext = os.path.splitext(fname)[1].lower()

        if ext in {".pdf", ".png", ".jpg", ".jpeg"}:
            # unified PDF + image OCR
            docs.extend(load_docs_with_ocr(path))

        elif ext in {".docx"}:
            # extract from Word
            docx = DocxDocument(path)
            text = "\n".join(p.text for p in docx.paragraphs if p.text)
            if text.strip():
                docs.append(Document(page_content=text))

        elif ext == ".txt":
            # plain text
            try:
                for d in TextLoader(path).load():
                    if d.page_content.strip():
                        docs.append(d)
            except Exception:
                continue
        else:
            print("Np document")        
        # skip any other file types

    return docs
